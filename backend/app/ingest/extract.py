"""Belge metni çıkarımı: PDF, DOCX, TXT/MD.

Çıktı, metnin tamamı ile birlikte sayfa sınırlarını (karakter aralığı) taşır;
böylece parçalanan her madde hangi sayfada geçtiğini bilir ve arayüzde
"sayfa 4, Madde 12" biçiminde atıf verilebilir.
"""

from __future__ import annotations

import re
import unicodedata
from dataclasses import dataclass, field
from pathlib import Path


@dataclass
class PageSpan:
    page: int
    start: int
    end: int


@dataclass
class ExtractedDocument:
    text: str
    pages: list[PageSpan] = field(default_factory=list)

    @property
    def page_count(self) -> int:
        return len(self.pages)

    def page_of(self, char_index: int) -> int | None:
        for span in self.pages:
            if span.start <= char_index < span.end:
                return span.page
        return self.pages[-1].page if self.pages else None


_WS = re.compile(r"[ \t   ]+")
_MULTI_NL = re.compile(r"\n{3,}")
# PDF'lerde satır sonu tirelemesi: "sözleş-\nme" -> "sözleşme"
_HYPHEN_BREAK = re.compile(r"(\w)-\n(\w)")


def clean_text(raw: str) -> str:
    text = unicodedata.normalize("NFKC", raw)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = _HYPHEN_BREAK.sub(r"\1\2", text)
    text = _WS.sub(" ", text)
    text = "\n".join(line.strip() for line in text.split("\n"))
    text = _MULTI_NL.sub("\n\n", text)
    return text.strip()


def _extract_pdf(path: Path) -> ExtractedDocument:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    chunks: list[str] = []
    pages: list[PageSpan] = []
    cursor = 0
    for index, page in enumerate(reader.pages, start=1):
        page_text = clean_text(page.extract_text() or "")
        block = page_text + "\n\n"
        chunks.append(block)
        pages.append(PageSpan(page=index, start=cursor, end=cursor + len(block)))
        cursor += len(block)
    return ExtractedDocument(text="".join(chunks).strip(), pages=pages)


def _extract_docx(path: Path) -> ExtractedDocument:
    import docx

    document = docx.Document(str(path))
    parts: list[str] = [p.text for p in document.paragraphs]

    # Tablolar sözleşmelerde ücret/teslim planı taşır; satır bazında düzleştirilir.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    text = clean_text("\n".join(parts))
    return ExtractedDocument(text=text, pages=[PageSpan(page=1, start=0, end=len(text))])


def _extract_plain(path: Path) -> ExtractedDocument:
    text = clean_text(path.read_text(encoding="utf-8", errors="replace"))
    return ExtractedDocument(text=text, pages=[PageSpan(page=1, start=0, end=len(text))])


def extract(path: Path) -> ExtractedDocument:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        doc = _extract_pdf(path)
    elif suffix == ".docx":
        doc = _extract_docx(path)
    elif suffix in (".txt", ".md"):
        doc = _extract_plain(path)
    else:
        raise ValueError(f"Desteklenmeyen dosya türü: {suffix}")

    if not doc.text.strip():
        raise ValueError(
            "Belgeden metin çıkarılamadı. Dosya taranmış (görüntü tabanlı) bir PDF olabilir; "
            "OCR uygulanmış bir sürüm yükleyin."
        )
    return doc
